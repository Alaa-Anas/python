from docx import Document # creating a document
from docx.shared import Inches # controlling the image size

document = Document()

document.add_picture( # add picture
  'jqz2YF0.jpg',
  width=Inches(2.0)
)
name = input('what is your name?')# varibles
phone = input('what is your phone number?') # varibles
email = input('what is your email?') # varibles
# putting varibles inside a paragraph
document.add_paragraph( name + "|" + phone + "|" + email + "|")

document.add_heading('reading') # add heading

# add paragraph in a varibles
p = document.add_paragraph() 
university = input("what university did you go to?")
coursenum = input("how many courses did you take?")
totalfees= input("how much did you pay?")

p.add_run(university + " ").bold = True
p.add_run(coursenum + " " + totalfees).italic = True

your_knowledge = input("what knowledge doyou have from your"+ university)

p.add_run()

document.save("cv.docx") # closing the document