from docx import Document # creating a document
import pyttsx3

def speak(text):
    pyttsx3.speak(text)
speak('hello ')

document = Document()

document.add_heading('at an interview') # add heading

# add paragraph in a varibles
p = document.add_paragraph() 

university = input("what university did you go to?")

coursenum = input("how many courses did you take?")
totalfees= input("how much did you pay?")

p.add_run(university + " ").bold = True
p.add_run(coursenum + " " + totalfees).italic = True

your_knowledge = input("what knowledge do you have from your" + " " + university)
while True: # using while & if to tell about more knowledge
    more_knowledge = input('do you have more knowledge, reply yes or no ')
    if more_knowledge == 'yes':
        university = input("what university did you go to?")
        coursenum = input("how many courses did you take?")
        totalfees= input("how much did you pay?")

        p.add_run(university + " ").bold = True
        p.add_run(coursenum + " " + totalfees).italic = True
    else:
      break

p.add_run()

document.save("cv.docx") # closing the document